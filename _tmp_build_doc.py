# -*- coding: utf-8 -*-
import copy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

SRC = 'Huong_dan_su_dung.docx'
OUT = 'Huong_dan_su_dung.docx'
SHOT_DIR = r'C:\Users\longn\AppData\Local\Temp\claude_screens'

CAPTION_SIZE = Pt(9)
CAPTION_COLOR = RGBColor(0x80, 0x80, 0x80)
IMG_WIDTH = Inches(2.4)

doc = Document(SRC)
paras = doc.paragraphs

placeholders = [
    ('Trang chủ học viên', 'home.png'),
    ('Lớp đang học', 'classes.png'),
    ('Lịch học', 'schedule.png'),
    ('Điểm danh cá nhân', 'attendance.png'),
    ('Học phí học viên', 'fees.png'),
    ('Lịch sử thanh toán', 'payments.png'),
    ('Tài liệu', 'documents.png'),
    ('Thông báo', 'notifications.png'),
    ('Hồ sơ cá nhân', 'profile.png'),
    ('Mục Tài khoản', 'more.png'),
]

def set_caption_run(run, text):
    run.text = text
    run.italic = True
    run.font.size = CAPTION_SIZE
    run.font.color.rgb = CAPTION_COLOR

placeholder_idx = 0
i = 0
while i < len(paras):
    p = paras[i]
    if p.text.strip().startswith('[Chua co anh minh hoa') or p.text.strip().startswith('[Chưa có ảnh minh họa'):
        caption_text, img_file = placeholders[placeholder_idx]
        placeholder_idx += 1
        # Clear existing runs
        for r in list(p.runs):
            r.text = ''
        # Remove all run elements then add fresh picture run
        for r in list(p.runs):
            r._element.getparent().remove(r._element)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(f'{SHOT_DIR}\\{img_file}', width=IMG_WIDTH)
        # Caption paragraph is the very next paragraph (was blank/placeholder text originally single-purpose)
        # Insert a new caption paragraph right after
        new_p_elm = OxmlElement('w:p')
        p._p.addnext(new_p_elm)
        cap_para = Paragraph(new_p_elm, p._parent)
        cap_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap_run = cap_para.add_run()
        set_caption_run(cap_run, caption_text)
    i += 1

assert placeholder_idx == 10, f'Expected to replace 10 placeholders, replaced {placeholder_idx}'

doc.save(OUT)
print('Saved with', placeholder_idx, 'images inserted')
